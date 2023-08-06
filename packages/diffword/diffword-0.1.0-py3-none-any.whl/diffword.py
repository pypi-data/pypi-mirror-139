from difflib import ndiff
from textwrap import wrap

import click
import regex as re
from docx import Document

DIFF_COLORS = {"^": "yellow", "-": "red", "+": "green"}


def diff(lines1, lines2):
    prev_line = ""
    for line in ndiff(lines1, lines2):
        if not (line.startswith("  ") or line.strip() == ""):
            line = line.strip()
            if line.startswith("? "):
                c_prev_line = []
                prev_index = 0
                for m in re.finditer(r"(\^|\+|-)", line):
                    start, end = m.span()
                    if prev_index < start:
                        c_prev_line.append(prev_line[prev_index:start])
                    c_prev_line += click.style(
                        prev_line[start:end], fg=DIFF_COLORS[line[start]], bold=True
                    )
                    prev_index = end
                c_prev_line += prev_line[prev_index:]
                prev_line = "".join(c_prev_line)
        if not (
            prev_line.startswith("  ")
            or prev_line.strip() == ""
            or prev_line.startswith("? ")
        ):
            yield prev_line
        prev_line = line
    if not (
        prev_line.startswith("  ")
        or prev_line.strip() == ""
        or prev_line.startswith("? ")
    ):
        yield prev_line


@click.command()
@click.argument("doc1")
@click.argument("doc2")
def main(doc1, doc2):
    click.echo(f"- {doc1}\n+ {doc2}")

    d1 = Document(doc1)
    d2 = Document(doc2)

    d = tuple(
        diff(
            [
                l.strip()
                for p in d1.paragraphs
                for l in p.text.splitlines()
                if l.strip() != ""
            ],
            [
                l.strip()
                for p in d2.paragraphs
                for l in p.text.splitlines()
                if l.strip() != ""
            ],
        )
    )

    if len(d) > 0:
        click.echo("\n\t".join(("Changes outside of tables:",) + d), color=True)

    if len(d1.tables) != len(d2.tables):
        click.echo(
            f"{doc1} and {doc2} do not have the same number of tables "
            f"({len(d1.tables)} and {len(d2.tables)} respectively)."
        )
    n_tables = min(len(d1.tables), len(d2.tables))

    diffs = set()
    for table_n in range(n_tables):
        t1 = d1.tables[table_n]
        t2 = d2.tables[table_n]
        if len(t1.rows) != len(t2.rows):
            click.echo(
                f"Table n°{table_n} from {doc1} and {doc2} do not have the same number of lines "
                f"({len(t1.rows)} and {len(t2.rows)} respectively)."
            )
        n_rows = min(len(t1.rows), len(t2.rows))

        for row_n in range(n_rows):
            r1 = t1.rows[row_n]
            r2 = t2.rows[row_n]
            if len(r1.cells) != len(r2.cells):
                click.echo(
                    f"Row n°{row_n} of table n°{table_n} from {doc1} and {doc2} do not have the same number of cells "
                    f"({len(r1.cells)} and {len(r2.cells)} respectively)."
                )
            n_cells = min(len(r1.cells), len(r2.cells))

            for cell_n in range(n_cells):
                c1 = r1.cells[cell_n]
                c2 = r2.cells[cell_n]
                if c1.text.strip().lower() != c2.text.strip().lower():
                    d = tuple(
                        diff(
                            [
                                l1
                                for l2 in c1.text.strip().splitlines()
                                for l1 in wrap(l2, width=120)
                            ],
                            [
                                l1
                                for l2 in c2.text.strip().splitlines()
                                for l1 in wrap(l2, width=120)
                            ],
                        )
                    )

                    if len(d) > 0 and (f"t{table_n}", d) not in diffs:
                        diffs.add((f"t{table_n}", d))
                        click.echo(
                            "\n\t".join(
                                (
                                    f"Table n°{table_n}, row n°{row_n}, column n°{cell_n}:",
                                )
                                + d
                            ),
                            color=True,
                        )
